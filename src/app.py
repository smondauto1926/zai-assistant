# Le librerie che sto importando servono a costruire un'applicazione web 
# con FastAPI per caricare, trascrivere e correggere file audio, 
# quindi generare un documento .docx con il testo formattato.

# Librerie di FastAPI (API web)
# FastAPI → Framework per creare API web veloci e asincrone.
# File, UploadFile → Gestiscono il caricamento di file tramite API.
# Request → Permette di accedere alle richieste HTTP.
# Jinja2Templates → Usa il motore di template Jinja2 per generare pagine HTML dinamiche.
# StaticFiles → Serve file statici (CSS, immagini, JavaScript).
# JSONResponse → Restituisce risposte in formato JSON.
# 🔹 Uso: Creare un server che riceve file audio, li elabora e restituisce il risultato.
from fastapi import FastAPI, File, UploadFile, Request
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import JSONResponse

#  Librerie per generare documenti Word
# docx.Document → Crea e modifica documenti Word (.docx).
# docx.shared.Pt, RGBColor → Imposta dimensione e colore del testo.
# docx.enum.text.WD_ALIGN_PARAGRAPH → Gestisce l'allineamento del testo.
# 🔹 Uso: Genera un documento Word con la trascrizione formattata.
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
 
# torch: Libreria di deep learning usata da modelli come Whisper (di OpenAI).
import torch

# whisper: Modello di OpenAI per la trascrizione automatica dell'audio in testo.
# 🔹 Uso: Carica Whisper con PyTorch per trascrivere i file audio.
import whisper

# Operazioni sui file (es. spostare, copiare, eliminare).
import shutil
# tempfile → Crea file e cartelle temporanee per salvare gli audio.
import tempfile
# json → Gestisce il formato JSON per la comunicazione tra API.
import json
# os → Interagisce con il filesystem (es. leggere/scrivere file).
import os
# Libreria per eseguire l'applicazione
import uvicorn  

# Librerie per correzione grammaticale e AI
# interfaccia per LanguageTool, che corregge errori grammaticali e ortografici.
import language_tool_python

# Client per usare le API di OpenAI (es. ChatGPT o modelli di completamento testi).
import openai

# Inizializza un'app FastAPI, che fungerà da server web.
# Questa istanza app sarà usata per definire API, gestire richieste e rispondere ai client.
app = FastAPI()

# Monta la cartella static/ nella route /static.
# Qualsiasi file in static/ sarà accessibile via URL, ad esempio:
# http://localhost:8000/static/styles.css
# Questo è utile per servire file CSS, immagini, JavaScript o documenti generati.
app.mount("/static", StaticFiles(directory="static"), name="static")

# Configura Jinja2, un motore di template per generare pagine HTML dinamiche.
# I file HTML devono trovarsi nella cartella templates/.
# Può essere usato per passare dati a una pagina HTML.
templates = Jinja2Templates(directory="templates")

# caricamento del modello whisper "tiny" per la trascrizione di un aduo
device = "cpu"  # Change to "cuda" if using GPU
model = whisper.load_model("tiny", device = device)
model = model.to(torch.float32)  

# strumento per correggere la grammatica italiana "it"
tool = language_tool_python.LanguageToolPublicAPI("it")

# Caricare le credenziali di openai per chiamare i modelli di GPT
with open('openai.json', 'r') as file:
    openai_credentials = json.load(file)

API_KEY = openai_credentials['API_KEY']
openai.api_key = API_KEY

# La seguente è una funzione che ha il compito di trascrivere un audio utilizzato il 
# modello gpt-4
def improve_transcript(raw_transcript):

    prompt = f"""Migliora questa trascrizione mantenendo il significato originale ma correggendo solamente:
    - Errori grammaticali
    - Punteggiatura

    Nel trascritto ci sono parole chiavi tali che devono essere solo corrette grammaticalmente e ortograficamente:
    - Caratteristiche costruttive dei fabbricati
    - Hot work
    - Ciclo produttivo
    
    Trascrizione originale:
    {raw_transcript}
    
    """

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}]
    )
    
    return response["choices"][0]["message"]["content"]


# inizializza l'applicazione utilizzando la template definita nel file templates/index.html
@app.get("/")
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request, "generated_file": None})

@app.post("/transcribe/")
async def transcribe_audio(request: Request, file: UploadFile = File(...)):
    # Salva il file temporaneamente
    with tempfile.NamedTemporaryFile(delete=False, suffix=".mp3") as temp_audio:
        shutil.copyfileobj(file.file, temp_audio)
        temp_audio_path = temp_audio.name

    try:
        # passaggio 1) qui è dove viene eseguita la trascrizione dell'audio file
        result = model.transcribe(temp_audio_path)
        raw_transcribed_text = result["text"]

        # passaggio 2) qui correggiamo l'ortografia e la grammatica del testo italiano
        improved_transcript = improve_transcript(raw_transcribed_text)

        # passaggio 3) qui generiamo il report finale utilizzando le istruzioni seguenti:
        analysis_prompt = f"""
        Dato un testo di input, crea un report che abbia come capoverso le seguenti parole: 
        - Caratteristiche costruttive dei fabbricati
        - Hot work
        - Sistemi di spegnimento automatico
        - Rilevatori di fumo
        - Fotovoltaico
        - Ciclo produttivo

        Riporta esclusivamente il contenuto testuale senza aggiungere spiegazioni, dettagli o interpretazioni.        
        Alla fine, crea un paragrafo conclusivo di riepilogo.

        Se non trovi contenuto per le parole sopracitate, riporta: Informazione non presente 
        
        
        Questo è il testo in input da analizzare:\n\n{improved_transcript}"""
        
        # qui chiamiamo il modello gpt-4-turbo con le istruzioni sopracitate.
        analysis_response = openai.ChatCompletion.create(
            model="gpt-4-turbo",
            messages=[{"role": "user", "content": analysis_prompt}]
        )
        ai_analysis = analysis_response["choices"][0]["message"]["content"]

        # creazione del documento word
        doc_path = "static/generated_summary.docx"
        document = Document()
        
        # 1. Aggiungi titolo
        title = document.add_heading("Report AMFU Italia", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        # 2. Aggiungi paragrafo Trascritto
        document.add_heading("Trascritto", level=2)
        document.add_paragraph(improved_transcript)
        
        # Aggiungi separatore
        document.add_paragraph("_" * 50)
        
        # 3. Aggiungi paragrafo con il report finale
        document.add_heading("REPORT", level=2)
        analysis_para = document.add_paragraph()
        analysis_para.add_run(":\n\n").italic = True
        document.add_paragraph(ai_analysis)

        # Salva il documento
        document.save(doc_path)

        os.unlink(temp_audio_path)

        return JSONResponse(content={"generated_file": "/static/generated_summary.docx"})

    except Exception as e:
        # Clean up temporary file in case of error
        if os.path.exists(temp_audio_path):
            os.unlink(temp_audio_path)
        raise e

# Esegui l'applicazione tramite uvicorn in localhost/sul computer. 
if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
